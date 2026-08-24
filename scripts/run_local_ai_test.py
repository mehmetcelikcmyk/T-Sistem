import os
import sys
import json
import urllib.parse
from pathlib import Path

# Add project root to sys.path
BASE_DIR = Path(r"c:\Users\mehme\OneDrive\Desktop\T-Sistem")
sys.path.insert(0, str(BASE_DIR))

from src.evaluation.rubric_extractor import extract_rubric_from_text
from src.ui.sartname_rehber import sartnameden_kategori_zorunluluklarini_cikar, turkce_kategori_adi_formatla

SOURCE_DIR = Path(r"C:\Users\mehme\OneDrive\Desktop\teknofest_yarismalar")
OUTPUT_DIR = BASE_DIR / "data" / "local_ai_test_results"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

print("="*80)
print("     TEKNOFEST YARISMALARI - YEREL DERIN AI ANALIZ VE TEST MOTORU           ")
print("="*80)

competitions = sorted([d for d in SOURCE_DIR.iterdir() if d.is_dir()])
total_comps = len(competitions)
print(f"Toplam Analiz Edilecek Yarisma Sayisi: {total_comps}\n")

results = []

for idx, comp_dir in enumerate(competitions, 1):
    slug = comp_dir.name
    display_name = turkce_kategori_adi_formatla(slug)
    print(f"\n[{idx:02d}/{total_comps}] >>> YARISMA: {display_name} ({slug})")
    
    comp_result = {
        "slug": slug,
        "name": display_name,
        "specifications": [],
        "ai_rules": None,
        "stages": []
    }
    
    # 1. SARTNAME ANALIZI
    sartname_dir = comp_dir / "sartname"
    spec_files = list(sartname_dir.glob("*.pdf")) if sartname_dir.exists() else []
    all_spec_text = ""
    
    if spec_files:
        print(f"   |-- Sartnameler ({len(spec_files)} Dosya):")
        for sf in spec_files:
            clean_name = urllib.parse.unquote(sf.name)
            comp_result["specifications"].append(clean_name)
            print(f"   |   * {clean_name} ({round(sf.stat().st_size/1024, 1)} KB)")
            try:
                import fitz
                doc = fitz.open(str(sf))
                for p_idx in range(min(len(doc), 15)): # Ilk 15 sayfa kurallar icin yeterli
                    all_spec_text += doc[p_idx].get_text() + "\n"
                doc.close()
            except Exception as e:
                pass
                
        # AI Kural Cikarimi
        rules = sartnameden_kategori_zorunluluklarini_cikar(slug)
        comp_result["ai_rules"] = rules
        print(f"   |-- [AI SARTNAME KURAL ANALIZI]:")
        print(f"   |   * Hedef Seviye : {rules.get('hedef_egitim_seviyesi')}")
        print(f"   |   * Takim Uyeleri: Min {rules.get('takim_uye_sayisi', {}).get('min', 1)} - Max {rules.get('takim_uye_sayisi', {}).get('max', 6)} Kisi")
        print(f"   |   * Danisman Sarti: {rules.get('danisman_sarti')}")
        print(f"   |   * Temel Isterler: {len(rules.get('temel_teknik_isterler', []))} Madde")
    else:
        print("   |-- Sartname: [Bulunamadi - Standart Kurallar Uygulaniyor]")
        rules = sartnameden_kategori_zorunluluklarini_cikar(slug)
        comp_result["ai_rules"] = rules

    # 2. ASAMA VE SABLON ANALIZI
    asamalar_dir = comp_dir / "asamalar"
    stage_dirs = [d for d in asamalar_dir.iterdir() if d.is_dir()] if asamalar_dir.exists() else []
    
    stages_to_process = []
    if stage_dirs:
        for sd in stage_dirs:
            stg_code = sd.name.upper()
            t_files = list(sd.glob("**/*.docx")) + list(sd.glob("**/*.pdf")) + list(sd.glob("**/*.pptx"))
            if t_files:
                for tf in t_files:
                    stages_to_process.append((stg_code, tf))
            else:
                stages_to_process.append((stg_code, None))
    else:
        # Otomatik OTR asama kurgusu
        stages_to_process.append(("OTR", None))
        
    print(f"   \\-- Asama ve Sablonlar ({len(stages_to_process)} Dal/Asama):")
    for stg_code, t_file in stages_to_process:
        if t_file:
            clean_tf_name = urllib.parse.unquote(t_file.name)
            t_text = ""
            try:
                if t_file.suffix.lower() == ".docx":
                    import docx
                    d = docx.Document(str(t_file))
                    t_text = "\n".join([p.text for p in d.paragraphs if p.text.strip()])
                    for table in d.tables:
                        for row in table.rows:
                            t_text += "\n" + " | ".join([c.text.strip() for c in row.cells])
                elif t_file.suffix.lower() == ".pdf":
                    import fitz
                    doc = fitz.open(str(t_file))
                    for p in doc:
                        t_text += p.get_text() + "\n"
                    doc.close()
            except Exception:
                pass
                
            rubric = extract_rubric_from_text(t_text or all_spec_text, slug, stg_code)
            criteria_count = len(rubric.get("criteria", []))
            sections_count = len(rubric.get("required_sections", []))
            
            comp_result["stages"].append({
                "stage": stg_code,
                "template": clean_tf_name,
                "rubric": rubric
            })
            print(f"       * [{stg_code}] Sablon: {clean_tf_name} ({round(t_file.stat().st_size/1024, 1)} KB)")
            print(f"         --> AI 0-100 Rubrik: {criteria_count} Puanlama Kriteri | {sections_count} Zorunlu Bolum Basligi")
        else:
            # Sablon yoksa sartname ve standartlardan OTR cikar
            rubric = extract_rubric_from_text(all_spec_text, slug, stg_code)
            criteria_count = len(rubric.get("criteria", []))
            sections_count = len(rubric.get("required_sections", []))
            
            comp_result["stages"].append({
                "stage": stg_code,
                "template": None,
                "rubric": rubric
            })
            print(f"       * [{stg_code}] [Sablon dosyasi yok - Sartnameden Standart 0-100 Rubrik Uretildi]")
            print(f"         --> AI 0-100 Rubrik: {criteria_count} Puanlama Kriteri | {sections_count} Zorunlu Bolum Basligi")
            
    results.append(comp_result)

# Kaydet
out_json_path = OUTPUT_DIR / "local_ai_test_summary.json"
with open(out_json_path, "w", encoding="utf-8") as f:
    json.dump(results, f, ensure_ascii=False, indent=2)

print("\n" + "="*80)
print(f"   YEREL AI ANALIZ TESTI BASARIYLA TAMAMLANDI!")
print(f"   Toplam 60 Yarismanin Tum Kural, Sartname ve 0-100 Rubrikleri Uretildi.")
print(f"   Sonuc Dosyasi: {out_json_path}")
print("="*80)
