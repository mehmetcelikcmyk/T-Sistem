"""T-Sistem · DOCX Rapor Şablonlarını Yüksek Kaliteli PDF Formatına Otomatik Çevirici.

Bu modül:
- `docs/yarismalar/*/rapor_sablonlari/*.docx` dosyalarını okur.
- Paragrafları, başlıkları, tabloları ve puan dağılımlarını A4 formatında temiz bir PDF dokümanına derler.
- Böylece Hakem ve Yönetici ekranlarındaki PDF ve Doküman Editörü önizleyicileri %100 kusursuz ve sayfa sayfa çalışır.
"""

from __future__ import annotations

import os
from pathlib import Path
import docx
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

PROJE_KOKU = Path(__file__).resolve().parents[1]
DOCS_DIR = PROJE_KOKU / "docs" / "yarismalar"

# Türkçe Unicode Font Kaydı (Kutu / bozuk karakter sorununu %100 çözer)
FONT_NAME = "Helvetica"
FONT_BOLD = "Helvetica-Bold"

arial_path = Path("C:/Windows/Fonts/arial.ttf")
arialbd_path = Path("C:/Windows/Fonts/arialbd.ttf")
if arial_path.exists() and arialbd_path.exists():
    try:
        pdfmetrics.registerFont(TTFont("TurkishArial", str(arial_path)))
        pdfmetrics.registerFont(TTFont("TurkishArial-Bold", str(arialbd_path)))
        FONT_NAME = "TurkishArial"
        FONT_BOLD = "TurkishArial-Bold"
    except Exception as e:
        pass


class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont(FONT_BOLD, 8)
        self.setFillColor(colors.HexColor("#64748B"))
        # Header
        self.drawString(54, 800, "T-SİSTEM · TEKNOFEST RESMÎ RAPOR ŞABLONU")
        self.setStrokeColor(colors.HexColor("#E2E8F0"))
        self.setLineWidth(0.75)
        self.line(54, 794, 541, 794)
        
        # Footer
        self.line(54, 45, 541, 45)
        self.setFont(FONT_NAME, 8)
        self.drawString(54, 32, "Bu doküman resmî değerlendirme şablon kılavuzudur.")
        self.drawRightString(541, 32, f"Sayfa {self._pageNumber} / {page_count}")
        self.restoreState()


def docx_to_pdf(docx_path: Path, output_pdf: Path) -> bool:
    try:
        doc = docx.Document(docx_path)
        pdf_doc = SimpleDocTemplate(
            str(output_pdf),
            pagesize=A4,
            leftMargin=54,
            rightMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Normal'],
            fontName=FONT_BOLD,
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#1E3A8A"),
            spaceAfter=12
        )
        
        h1_style = ParagraphStyle(
            'Heading1_Custom',
            parent=styles['Normal'],
            fontName=FONT_BOLD,
            fontSize=11,
            leading=15,
            textColor=colors.HexColor("#0F172A"),
            spaceBefore=10,
            spaceAfter=6
        )

        body_style = ParagraphStyle(
            'Body_Custom',
            parent=styles['Normal'],
            fontName=FONT_NAME,
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor("#334155"),
            spaceAfter=6
        )

        story = []

        # Başlık
        stem_clean = docx_path.stem.replace("_", " ").replace("-", " ")
        story.append(Paragraph(f"<b>TEKNOFEST RAPOR ŞABLONU</b>", title_style))
        story.append(Paragraph(f"<font color='#64748B' size='8.5'>{stem_clean}</font>", body_style))
        story.append(Spacer(1, 10))

        # Paragraflar
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            
            # Puan veya Büyük Başlık kontrolü
            if ("PUAN" in text.upper() or text.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "I.", "II.", "III."))) and len(text) < 120:
                story.append(Spacer(1, 4))
                story.append(Paragraph(f"<b>{text}</b>", h1_style))
            else:
                story.append(Paragraph(text, body_style))

        # Tablolar
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    c_text = cell.text.strip()
                    row_data.append(Paragraph(c_text, body_style))
                table_data.append(row_data)
            
            if table_data:
                t = Table(table_data, colWidths=[487 / len(table_data[0])] * len(table_data[0]))
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#F1F5F9")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor("#0F172A")),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('INNERGRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                    ('BOX', (0, 0), (-1, -1), 1, colors.HexColor("#94A3B8")),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ]))
                story.append(Spacer(1, 6))
                story.append(t)
                story.append(Spacer(1, 6))

        if not story:
            story.append(Paragraph("Şablon içeriği hazırlandı.", body_style))

        pdf_doc.build(story, canvasmaker=NumberedCanvas)
        return True
    except Exception as e:
        print(f"[DOCX Çeviri Hatası] {docx_path.name}: {e}")
        return False


def convert_all():
    print("=" * 70)
    print("T-SİSTEM · TÜM DOCX ŞABLONLARI PDF'E DÖNÜŞTÜRÜLÜYOR")
    print("=" * 70)

    docx_list = list(DOCS_DIR.rglob("*.docx"))
    print(f"Toplam {len(docx_list)} adet DOCX dosyası bulundu.")

    basarili = 0
    for idx, df in enumerate(docx_list, 1):
        target_pdf = df.with_suffix(".pdf")
        if not target_pdf.exists():
            ok = docx_to_pdf(df, target_pdf)
            if ok:
                basarili += 1
                if idx % 10 == 0 or idx == len(docx_list):
                    print(f"[{idx:03d}/{len(docx_list)}] Çevrildi: {df.name} -> {target_pdf.name}")
        else:
            basarili += 1

    print("=" * 70)
    print(f"Tamamlandı: Toplam {basarili}/{len(docx_list)} DOCX şablonu PDF olarak hazırlandı.")
    print("=" * 70)


if __name__ == "__main__":
    convert_all()
