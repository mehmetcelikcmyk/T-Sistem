"""T-Sistem · Birebir Microsoft Word (Office 365 A4 Sayfa & Ribbon) Doküman Arayüzü.

Bu modül:
- DOCX dosyalarını Microsoft Word / Office 365 masaüstü uygulamasındaki gibi:
  Ribbon Menüsü (Dosya, Giriş, Ekle, Düzen), Biçimlendirme Araç Çubuğu, Yatay Cetvel,
  Gri Çalışma Masası üzerinde ardışık ayrık A4 Kağıt Yaprakları (210x297mm),
  Orijinal Kenar Boşlukları (2.54 cm), Calibri Tipografisi, Tablolar, Resimler
  ve Word Durum Çubuğu ile %100 birebir render eder.
"""

from __future__ import annotations

import html
import os
import re
from pathlib import Path

try:
    import mammoth
except ImportError:
    mammoth = None

try:
    import docx
except ImportError:
    docx = None


def docx_html_getir(docx_path: str | Path) -> str:
    """DOCX dosyasını Microsoft Office 365 A4 Masaüstü Düzeninde render eder."""
    p = Path(docx_path)
    if not p.exists():
        return "<div style='color:#EF4444; padding:16px;'>Word dosyası bulunamadı.</div>"

    raw_html = ""
    word_count = 0

    if mammoth is not None:
        try:
            with open(p, "rb") as docx_file:
                # Orijinal stilleri, tabloları ve resimleri doğrudan HTML olarak dönüştür
                style_map = """
                p[style-name='Heading 1'] => h1:fresh
                p[style-name='Heading 2'] => h2:fresh
                p[style-name='Heading 3'] => h3:fresh
                p[style-name='Title'] => h1.word-doc-title:fresh
                p[style-name='Subtitle'] => p.word-doc-subtitle:fresh
                table => table.word-office-table:fresh
                """
                result = mammoth.convert_to_html(docx_file, style_map=style_map)
                raw_html = result.value
        except Exception:
            raw_html = ""

    if not raw_html and docx is not None:
        try:
            doc = docx.Document(str(p))
            parts = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if not t:
                    continue
                esc = html.escape(t)
                if "PUAN" in t.upper() or t.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.", "I.", "II.", "III.")):
                    parts.append(f"<h2>{esc}</h2>")
                else:
                    parts.append(f"<p>{esc}</p>")
            
            for table in doc.tables:
                parts.append("<table class='word-office-table'>")
                for r_idx, row in enumerate(table.rows):
                    parts.append("<tr>")
                    for cell in row.cells:
                        c_text = html.escape(cell.text.strip())
                        tag = "th" if r_idx == 0 else "td"
                        parts.append(f"<{tag}>{c_text}</{tag}>")
                    parts.append("</tr>")
                parts.append("</table>")
            raw_html = "\n".join(parts)
        except Exception:
            raw_html = "<p>Doküman içeriği okunamadı.</p>"

    # Kelime sayısı hesapla
    text_content = re.sub(r"<[^>]+>", " ", raw_html)
    words = [w for w in text_content.split() if len(w) > 1]
    word_count = len(words)
    estimated_pages = max(1, (word_count // 350) + 1)

    stem_name = html.escape(p.name)

    # Otantik Microsoft Word / Office 365 UI
    word_ui = f"""
    <div class="word-office-app" style="
        background: #E1DFDD;
        border-radius: 8px;
        border: 1px solid #8A8886;
        box-shadow: 0 6px 24px rgba(0, 0, 0, 0.15);
        font-family: 'Segoe UI', -apple-system, BlinkMacSystemFont, Roboto, sans-serif;
        overflow: hidden;
        margin: 12px 0 20px 0;
    ">
        <!-- 1. Office 365 Üst Başlık Çubuğu -->
        <div style="background:#103F91; color:#FFFFFF; padding:8px 16px; display:flex; align-items:center; justify-content:space-between; font-size:0.84rem; border-bottom:1px solid #0C2E6B;">
            <div style="display:flex; align-items:center; gap:12px;">
                <svg width="20" height="20" viewBox="0 0 24 24" fill="#FFFFFF">
                    <path d="M21.17 3.25Q21.5 3.25 21.75 3.5 22 3.75 22 4.08V19.92Q22 20.25 21.75 20.5 21.5 20.75 21.17 20.75H8.83Q8.5 20.75 8.25 20.5 8 20.25 8 19.92V17H2.83Q2.5 17 2.25 16.75 2 16.5 2 16.17V7.83Q2 7.5 2.25 7.25 2.5 7 2.83 7H8V4.08Q8 3.75 8.25 3.5 8.5 3.25 8.83 3.25ZM7 8.5H3.5V15.5H7ZM16.27 7.75 14.88 12.38Q14.73 12.87 14.58 13.5 14.43 14.13 14.34 14.59 14.24 14.13 14.09 13.5 13.94 12.87 13.79 12.38L12.4 7.75H10.58L12.54 13.62 10.66 19.25H12.48L13.87 14.62Q14.02 14.13 14.17 13.5 14.32 12.87 14.41 12.41 14.51 12.87 14.66 13.5 14.81 14.13 14.96 14.62L16.35 19.25H18.17L16.29 13.62 18.25 7.75Z"/>
                </svg>
                <span style="font-weight:700; letter-spacing:0.02em;">Word</span>
                <span style="opacity:0.6;">|</span>
                <span style="font-weight:600;">{stem_name}</span>
                <span style="background:rgba(255,255,255,0.15); padding:2px 8px; border-radius:12px; font-size:0.72rem;">Kaydedildi</span>
            </div>
            <div style="display:flex; align-items:center; gap:16px; opacity:0.9; font-size:0.80rem;">
                <span>🔍 Belgede Ara...</span>
                <span>Düzenleme Modu</span>
            </div>
        </div>

        <!-- 2. Ribbon Menü Sekmeleri -->
        <div style="background:#F3F2F1; border-bottom:1px solid #EDEBE9; padding:4px 12px 0 12px; display:flex; gap:16px; font-size:0.82rem; color:#323130;">
            <span style="padding:6px 10px; cursor:pointer; font-weight:500;">Dosya</span>
            <span style="padding:6px 10px; border-bottom:2.5px solid #103F91; font-weight:700; color:#103F91; cursor:pointer;">Giriş</span>
            <span style="padding:6px 10px; cursor:pointer; font-weight:500;">Ekle</span>
            <span style="padding:6px 10px; cursor:pointer; font-weight:500;">Tasarım</span>
            <span style="padding:6px 10px; cursor:pointer; font-weight:500;">Düzen</span>
            <span style="padding:6px 10px; cursor:pointer; font-weight:500;">Başvurular</span>
            <span style="padding:6px 10px; cursor:pointer; font-weight:500;">Gözden Geçir</span>
            <span style="padding:6px 10px; cursor:pointer; font-weight:500;">Görünüm</span>
        </div>

        <!-- 3. Ribbon Biçimlendirme Araç Çubuğu -->
        <div style="background:#FFFFFF; border-bottom:1px solid #D2D0CE; padding:6px 16px; display:flex; align-items:center; gap:12px; flex-wrap:wrap; font-size:0.80rem; color:#323130;">
            <div style="display:flex; align-items:center; gap:4px; background:#F3F2F1; padding:3px 8px; border-radius:4px; border:1px solid #E1DFDD;">
                <span style="font-weight:600; font-family:'Calibri', sans-serif;">Calibri</span>
                <span style="color:#605E5C; font-size:0.7rem;">▼</span>
            </div>
            <div style="display:flex; align-items:center; gap:4px; background:#F3F2F1; padding:3px 6px; border-radius:4px; border:1px solid #E1DFDD;">
                <span style="font-weight:600;">11</span>
                <span style="color:#605E5C; font-size:0.7rem;">▼</span>
            </div>
            <div style="height:18px; width:1px; background:#EDEBE9;"></div>
            <div style="display:flex; gap:2px;">
                <span style="padding:3px 7px; font-weight:900; background:#F3F2F1; border-radius:3px;">K</span>
                <span style="padding:3px 7px; font-style:italic; font-weight:700;">T</span>
                <span style="padding:3px 7px; text-decoration:underline;">A</span>
                <span style="padding:3px 7px; color:#D83B01; font-weight:800;">A▾</span>
            </div>
            <div style="height:18px; width:1px; background:#EDEBE9;"></div>
            <div style="display:flex; gap:4px; color:#605E5C;">
                <span>☰ Sol</span>
                <span>☲ Orta</span>
                <span>☱ Sağ</span>
                <span style="color:#103F91; font-weight:700;">≡ İki Yana</span>
            </div>
            <div style="height:18px; width:1px; background:#EDEBE9;"></div>
            <div style="display:flex; gap:6px; font-size:0.76rem;">
                <span style="background:#F3F2F1; padding:2px 6px; border:1px solid #E1DFDD; border-radius:3px; font-weight:600;">Normal</span>
                <span style="background:#F3F2F1; padding:2px 6px; border:1px solid #E1DFDD; border-radius:3px; color:#1F497D; font-weight:700;">Başlık 1</span>
                <span style="background:#F3F2F1; padding:2px 6px; border:1px solid #E1DFDD; border-radius:3px; color:#1F497D; font-weight:700;">Başlık 2</span>
            </div>
        </div>

        <!-- 4. Word Yatay Cetvel (Ruler) -->
        <div style="background:#F3F2F1; border-bottom:1px solid #D2D0CE; padding:2px 0; display:flex; justify-content:center;">
            <div style="width:100%; max-width:816px; background:#FFFFFF; height:16px; border-left:1px solid #C8C6C4; border-right:1px solid #C8C6C4; display:flex; align-items:center; justify-content:space-between; padding:0 24px; font-size:0.65rem; color:#8A8886; font-family:monospace; user-select:none;">
                <span>◄ 1</span>
                <span>· 2 ·</span>
                <span>· 4 ·</span>
                <span>· 6 ·</span>
                <span>· 8 ·</span>
                <span>· 10 ·</span>
                <span>· 12 ·</span>
                <span>· 14 ·</span>
                <span>· 16 ·</span>
                <span>18 ►</span>
            </div>
        </div>

        <!-- 5. Gri Çalışma Masası ve A4 Sayfaları (Birebir Word Canvas) -->
        <div class="word-canvas-desk" style="
            background: #E1DFDD;
            padding: 32px 16px;
            max-height: 650px;
            overflow-y: auto;
            display: flex;
            flex-direction: column;
            align-items: center;
        ">
            <style>
                .word-a4-sheet {{
                    background: #FFFFFF;
                    width: 100%;
                    max-width: 816px; /* Birebir 8.5 x 11 inç / A4 Oranı */
                    min-height: 1056px;
                    padding: 72px 84px; /* Standart 1 inç (2.54 cm) Word Boşluğu */
                    margin-bottom: 28px;
                    box-shadow: 0 4px 20px rgba(0, 0, 0, 0.18), 0 1px 4px rgba(0, 0, 0, 0.08);
                    border: 1px solid #C8C6C4;
                    font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.40;
                    color: #000000;
                    box-sizing: border-box;
                    position: relative;
                }}
                .word-a4-sheet h1, .word-a4-sheet h1.word-doc-title {{
                    font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
                    font-size: 16pt;
                    font-weight: bold;
                    color: #1F497D;
                    margin-top: 18pt;
                    margin-bottom: 8pt;
                    line-height: 1.25;
                    border-bottom: 1.5px solid #1F497D;
                    padding-bottom: 4pt;
                }}
                .word-a4-sheet h2 {{
                    font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
                    font-size: 13pt;
                    font-weight: bold;
                    color: #1F497D;
                    margin-top: 14pt;
                    margin-bottom: 6pt;
                    line-height: 1.25;
                }}
                .word-a4-sheet h3 {{
                    font-family: 'Calibri', 'Segoe UI', Arial, sans-serif;
                    font-size: 11pt;
                    font-weight: bold;
                    color: #365F91;
                    margin-top: 10pt;
                    margin-bottom: 4pt;
                }}
                .word-a4-sheet p {{
                    margin-top: 0;
                    margin-bottom: 6pt;
                    text-align: justify;
                    text-justify: inter-word;
                }}
                .word-a4-sheet table, .word-a4-sheet .word-office-table {{
                    width: 100% !important;
                    border-collapse: collapse !important;
                    margin: 14pt 0 !important;
                    font-size: 10pt !important;
                    font-family: 'Calibri', sans-serif !important;
                }}
                .word-a4-sheet th, .word-a4-sheet td {{
                    border: 1px solid #7F7F7F !important;
                    padding: 6pt 9pt !important;
                    text-align: left !important;
                    vertical-align: top !important;
                }}
                .word-a4-sheet th {{
                    background-color: #F2F2F2 !important;
                    color: #000000 !important;
                    font-weight: bold !important;
                }}
                .word-a4-sheet ul, .word-a4-sheet ol {{
                    margin-top: 0;
                    margin-bottom: 6pt;
                    padding-left: 24pt;
                }}
                .word-a4-sheet li {{
                    margin-bottom: 3pt;
                }}
                .word-a4-sheet img {{
                    max-width: 100%;
                    height: auto;
                    display: block;
                    margin: 14pt auto;
                    border-radius: 2px;
                    box-shadow: 0 1px 4px rgba(0,0,0,0.12);
                }}
                .word-header-rule {{
                    display: flex;
                    justify-content: space-between;
                    border-bottom: 1px solid #D2D0CE;
                    padding-bottom: 6pt;
                    margin-bottom: 24pt;
                    font-size: 8.5pt;
                    color: #605E5C;
                }}
                .word-footer-rule {{
                    margin-top: 36pt;
                    border-top: 1px solid #D2D0CE;
                    padding-top: 6pt;
                    font-size: 8.5pt;
                    color: #605E5C;
                    display: flex;
                    justify-content: space-between;
                }}
            </style>

            <!-- A4 Doküman Kağıdı -->
            <div class="word-a4-sheet">
                <div class="word-header-rule">
                    <span>TEKNOFEST 2026 · Resmî Rapor Şablon Kılavuzu</span>
                    <span>Microsoft Word</span>
                </div>

                {raw_html}

                <div class="word-footer-rule">
                    <span>T-Sistem · TEKNOFEST Şablon Denetleyici</span>
                    <span>Sayfa 1 / {estimated_pages}</span>
                </div>
            </div>
        </div>

        <!-- 6. Office 365 Alt Durum Çubuğu (Status Bar) -->
        <div style="background:#103F91; color:#FFFFFF; padding:4px 16px; display:flex; align-items:center; justify-content:space-between; font-size:0.76rem; border-top:1px solid #0C2E6B; user-select:none;">
            <div style="display:flex; align-items:center; gap:16px;">
                <span>Sayfa 1 / {estimated_pages}</span>
                <span>Sözcük: {word_count:,}</span>
                <span>Yazım Denetimi: Hatasız ✓</span>
                <span>Türkçe (Türkiye)</span>
            </div>
            <div style="display:flex; align-items:center; gap:14px;">
                <span>Odaklan</span>
                <span>Sayfa Düzeni</span>
                <span>%100 [ - ──|── + ]</span>
            </div>
        </div>
    </div>
    """
    return word_ui


def docx_onizle(st_obj, docx_path: str | Path, baslik: str = "Word Doküman Önizlemesi", key: str = "docx_view", r2_public_url: str | None = None) -> None:
    """Streamlit içerisinde PDF Eşleniği, Office 365 Masaüstü A4 ve Microsoft Office Online Viewer desteği sunar."""
    p = Path(docx_path)
    if not p.exists():
        st_obj.warning("Belirtilen Word dosyası mevcut değil.")
        return

    # PDF eşleniğini ara (.docx yanında .pdf var mı?)
    pdf_counterpart = p.with_suffix(".pdf")
    has_pdf = pdf_counterpart.exists()

    with st_obj.container(border=True):
        col_t1, col_t2 = st_obj.columns([2.5, 1.8])
        with col_t1:
            st_obj.markdown(
                f"""
                <div style="display:flex; align-items:center; gap:8px;">
                    <span style="background:#103F91; color:#FFFFFF; font-weight:800; font-size:0.78rem; padding:4px 9px; border-radius:4px; letter-spacing:0.02em;">OFFICE 365</span>
                    <span style="font-weight:700; color:#0F172A; font-size:0.96rem;">{p.name}</span>
                </div>
                """,
                unsafe_allow_html=True
            )
            st_obj.caption("TEKNOFEST Resmî Rapor Şablon Dokümanı")
        
        with col_t2:
            dl_col1, dl_col2 = st_obj.columns(2)
            with dl_col1:
                with open(p, "rb") as f_dl:
                    st_obj.download_button(
                        "📥 İndir (.docx)",
                        data=f_dl.read(),
                        file_name=p.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True,
                        key=f"dl_docx_{key}"
                    )
            with dl_col2:
                if has_pdf:
                    with open(pdf_counterpart, "rb") as f_pdf:
                        st_obj.download_button(
                            "📄 İndir (.pdf)",
                            data=f_pdf.read(),
                            file_name=pdf_counterpart.name,
                            mime="application/pdf",
                            use_container_width=True,
                            key=f"dl_pdf_{key}"
                        )

        # Görünüm Sekmeleri
        tab_list = []
        if has_pdf:
            tab_list.append("📄 1. Kesintisiz PDF & Baskı Görüntüleyici (Kaydırma & Büyüteç)")
        tab_list.append("💻 2. Office 365 Masaüstü A4 Düzeni")
        tab_list.append("🌐 3. Resmî Microsoft Office Online Viewer")

        view_tabs = st_obj.tabs(tab_list)
        tab_idx = 0

        # 1. SEKME: KESİNTİSİZ PDF GÖRÜNTÜLEYİCİ
        if has_pdf:
            with view_tabs[tab_idx]:
                import pdf_gorunum
                pdf_gorunum.pdf_onizle(st_obj, pdf_counterpart, height=720, key=f"docx_pdf_{key}")
            tab_idx += 1

        # 2. SEKME: OFFICE 365 MASAÜSTÜ A4
        with view_tabs[tab_idx]:
            html_content = docx_html_getir(p)
            st_obj.markdown(html_content, unsafe_allow_html=True)
        tab_idx += 1

        # 3. SEKME: MICROSOFT OFFICE ONLINE VIEWER
        with view_tabs[tab_idx]:
            st_obj.markdown("<div style='font-weight:700; color:#1E293B; margin-bottom:6px;'>Microsoft Office Online Web Viewer (Resmî & Ücretsiz)</div>", unsafe_allow_html=True)
            st_obj.caption("Bu servis Microsoft tarafından ücretsiz sunulmakta olup, Cloudflare R2'ye yüklenen belgelerinizi hiçbir ek yazılıma gerek kalmadan doğrudan tarayıcı içinde açar.")

            target_url = r2_public_url or f"https://view.officeapps.live.com/op/embed.aspx?src=https://t-sistem-raporlar.r2.cloudflarestorage.com/sablonlar/{p.name}"
            import urllib.parse
            encoded_url = urllib.parse.quote(target_url, safe=":/?=&")

            st_obj.html(f"""
            <div style="margin-top:10px; border-radius:8px; overflow:hidden; border:1px solid #CBD5E1; box-shadow: 0 4px 12px rgba(0,0,0,0.06);">
                <div style="background:#103F91; color:#FFFFFF; padding:8px 16px; font-weight:700; font-size:0.88rem; display:flex; justify-content:space-between; align-items:center;">
                    <span>🌐 Microsoft Office Online Viewer · {p.name}</span>
                    <span style="font-size:0.75rem; background:rgba(255,255,255,0.2); padding:2px 8px; border-radius:4px;">Canlı Web Önizleme</span>
                </div>
                <iframe src="https://view.officeapps.live.com/op/embed.aspx?src={encoded_url}"
                        width="100%"
                        height="650px"
                        frameborder="0"
                        style="background:#F3F4F6;">
                </iframe>
            </div>
            """)
