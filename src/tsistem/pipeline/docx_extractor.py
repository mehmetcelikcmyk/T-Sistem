"""DOCX okuma katmanı.

Neden gerekli: TEKNOFEST rapor şablonlarını `.docx` olarak yayınlıyor
(ör. "SIHA_KTR_2026_TR_v2.docx"). Bu şablonlar bizim için altın değerinde:

  * Zorunlu başlıkların TAM yazımını içeriyor → `data/templates/*.json`
  * Word'ün stil bilgisi (Heading 1/2/3) sayesinde hiyerarşi kesin;
    PDF'te tipografiden tahmin etmek zorundayız, burada tahmin yok.
  * **Puan ağırlıkları başlıkta yazıyor** — "OTONOM GÖREVLER (25 Puan)".
    Bu, MVP 6'daki kriter değerlendirmesinin ağırlıklarını doğrudan veriyor.

GERÇEK ŞABLONLARDAN ÖĞRENİLENLER (2026 şablonları incelenerek):

  * Her şablon ana bölümleri aynı seviyede tutmuyor. Jet Motor DTR'de
    Heading 1 önsöz için (ŞEKİLLER, TABLOLAR, SİMGELER), ana bölümler
    Heading 2'de. Savaşan İHA'da ana bölümler Heading 1'de. Bu yüzden
    "ana bölüm seviyesi" sabit değil, veriden tespit ediliyor.
  * Bazı şablonlarda yönlendirme cümlesi yanlışlıkla başlık stiliyle
    biçimlendirilmiş (Robotaksi Özgün KTR: "Bu kısımda proje hakkında
    genel bilgiler verilmelidir." → Heading 1). Bunlar eleniyor.
  * Bazı gerçek bölümler hiç başlık stili almamış, sadece kalın yazılmış
    (Roket AHR: "Özet", "Giriş"). Bunları alıyoruz ama `detected_by="bold"`
    işaretiyle — insan gözden geçirsin diye.

Takımlar raporu genelde PDF yüklediği için ANALİZ yolu PDF üzerinden kalıyor;
docx yolu şablon okumak ve docx yüklenen raporu işlemek için.
"""

from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path

from ..models import DocumentMeta, ExtractedDocument, Heading, PageText
from .language import detect_language
from .section_parser import normalize

logger = logging.getLogger(__name__)

#: Word'ün başlık stilleri. Türkçe Office "Başlık 1" yazar.
#: "Numarasız Başlık" gibi türetilmiş stiller kasten eşleşmiyor — onlar
#: şablonlarda önsöz (şekil/tablo listesi) için kullanılıyor.
HEADING_STYLE_RE = re.compile(r"^(heading|başlık|baslik|title)\s*(\d*)\s*$", re.IGNORECASE)

NUMBERING_RE = re.compile(
    r"^\s*((?:\d+(?:\s*\.\s*\d+)*)|(?:[IVXLC]+)|(?:[A-ZÇĞİÖŞÜ]))\s*[\.\)\-–]\s+"
)

#: Şablonlar puan ağırlığını parantez içinde yazıyor ama biçim tutarsız:
#:   "(30 Puan)"  "(2.5 PUAN)"  "(10 puan)"
#:   "(ve RAPOR DÜZENİ 10 PUAN)"   <- Sağlıkta YZ PDR'de böyle
#: O yüzden parantez içinde "sayı + puan" ikilisini herhangi bir yerde arıyoruz.
POINTS_RE = re.compile(r"\([^()]*?(\d+(?:[.,]\d+)?)\s*puan[ıi]?[^()]*\)", re.IGNORECASE)

#: TEKNOFEST'in bazı şablonlarında (ör. Roket AHR) bölümler açıkça
#: işaretlenmiş:  "Bölüm Başlangıcı – IX.Uçuş Kontrol Bilgisayarı"
#: Bu, tipografi sezgilerinden KAT KAT güvenilir: şablonun kendi beyanı.
#: Hangi başlıkların ana bölüm olduğunu ve sınırlarını kesin veriyor.
SECTION_MARKER_RE = re.compile(
    r"^\s*bölüm\s+(başlangıcı|baslangici|sonu)\s*[–\-—:]\s*"
    r"([IVXLC]+|\d+)\s*[\.\)]?\s*(.*)$",
    re.IGNORECASE,
)

#: Şablonlardaki yönlendirme/talimat metinleri
PLACEHOLDER_HINTS = (
    "bu bölümde", "bu kısımda", "bu başlık altında", "açıklanmalıdır",
    "belirtilmelidir", "yazılmalıdır", "anlatılmalıdır", "verilmelidir",
    "doldurulacak", "eklenmelidir", "siliniz", "silinmelidir",
    "örnek:", "not:", "aşağıdaki tabloda", "aşağıdaki şablon",
    "puanlaması aşağıdaki", "gibi doldurulmalıdır",
)


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for block in iter(lambda: fh.read(1 << 20), b""):
            h.update(block)
    return h.hexdigest()


def looks_like_placeholder(text: str) -> bool:
    """Şablondaki yönlendirme metni mi, gerçek içerik/başlık mı?

    Gerçek şablonlarda bu cümleler bazen başlık stiliyle biçimlendirilmiş
    oluyor; bölüm sanılmamaları için başlık adaylarına da uygulanıyor.
    """
    low = text.lower()
    return any(h in low for h in PLACEHOLDER_HINTS)


def extract_points(text: str) -> tuple[str, float | None]:
    """'OTONOM GÖREVLER (25 Puan)' -> ('OTONOM GÖREVLER', 25.0)"""
    m = POINTS_RE.search(text)
    if not m:
        return text.strip(), None
    value = float(m.group(1).replace(",", "."))
    cleaned = POINTS_RE.sub("", text).strip(" \t-–—:")
    return cleaned or text.strip(), value


def _style_level(style_name: str | None) -> int | None:
    """'Heading 2' -> 2, 'Başlık 1' -> 1, 'Title' -> 0, gövde -> None."""
    if not style_name:
        return None
    m = HEADING_STYLE_RE.match(style_name.strip())
    if not m:
        return None
    if m.group(1).lower() == "title":
        return 0
    digits = m.group(2)
    return int(digits) if digits else 1


def _is_all_bold(paragraph) -> bool:
    runs = [r for r in paragraph.runs if (r.text or "").strip()]
    return bool(runs) and all(r.bold for r in runs)


def _plausible_bold_heading(text: str) -> bool:
    """Kalın metin gerçekten başlık olabilir mi?

    Gerçek şablonlarda kalın-tespit yöntemi çok gürültü üretiyordu:
    tablo hücreleri, parantez içinde kalmış puan açıklamaları
    ("(Her bir madde 1 puan toplam 10 puan)"), satır kırılması artıkları
    ("Puan)"). Bu filtre onları eliyor.
    """
    t = text.strip()
    if not (3 <= len(t) <= 90):
        return False
    words = t.split()
    if not (1 <= len(words) <= 9):
        return False
    if t.startswith(("(", "[", "•", "-", "–")):
        return False
    if t.endswith((")", ",", ";")) and not POINTS_RE.search(t):
        return False
    if looks_like_placeholder(t):
        return False
    letters = [c for c in t if c.isalpha()]
    if not letters:
        return False
    # Başlıklar ya BÜYÜK HARF ya Başlık Biçimi olur; cümle gibi olanı alma
    upper_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
    first_caps = sum(1 for w in words if w[:1].isupper()) / len(words)
    return upper_ratio > 0.6 or first_caps >= 0.75


def extract_docx(path: str | Path) -> tuple[ExtractedDocument, list[Heading]]:
    """DOCX'i metin + başlık listesi olarak çıkarır.

    Word'de sayfa kavramı akışkan olduğu için (sayfalar render sırasında
    oluşur) tek "sayfa" olarak döndürülür; karakter ofsetleri tutarlı kalır.

    Başlık tespiti iki kaynaktan: Word stili (kesin, `detected_by="style"`)
    ve kalın metin sezgisi (belirsiz, `detected_by="bold"`). İkincisi yalnız
    ilk filtreyi geçen satırlar için üretilir ve çağıran taraf isterse
    ayırt edebilir.
    """
    import docx  # tembel import — docx desteği opsiyonel

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX bulunamadı: {path}")

    document = docx.Document(str(path))
    parts: list[str] = []
    headings: list[Heading] = []
    cursor = 0

    def add_text(text: str) -> None:
        nonlocal cursor
        parts.append(text)
        cursor += len(text) + 2  # "\n\n" ayırıcı

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue

        try:
            style_name = para.style.name if para.style else None
        except Exception:
            style_name = None
        level = _style_level(style_name)
        bold = _is_all_bold(para)

        m = NUMBERING_RE.match(text)
        numbering = m.group(1).replace(" ", "") if m else None
        without_num = text[m.end():].strip() if m else text
        clean, points = extract_points(without_num)

        detected_by: str | None = None
        if level is not None:
            detected_by = "style"
        elif bold and _plausible_bold_heading(text):
            detected_by = "bold"

        if detected_by and clean:
            resolved_level = (
                level if level is not None
                else 1 + (numbering.count(".") if numbering else 0)
            )
            headings.append(
                Heading(
                    text=clean,
                    raw_text=text,
                    normalized=normalize(clean),
                    page_no=1,
                    char_start=cursor,
                    level=resolved_level,
                    numbering=numbering,
                    is_bold=bold,
                    detected_by=detected_by,
                    points=points,
                )
            )
        add_text(text)

    # Tablo içerikleri de metne dahil (TEKNOFEST şablonları tablo yoğun).
    # Tablo hücreleri başlık ADAYI olarak DEĞERLENDİRİLMİYOR — gerçek
    # şablonlarda bu, sahte bölüm üretmenin ana kaynağıydı.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if (c.text or "").strip()]
            if cells:
                add_text(" | ".join(cells))

    full_text = "\n\n".join(parts)
    lang, conf = detect_language(full_text)

    meta = DocumentMeta(
        file_name=path.name,
        file_sha256=_sha256(path),
        page_count=1,
        pdf_title=None,
        is_scanned=False,
        total_chars=len(full_text),
    )
    extracted = ExtractedDocument(
        meta=meta,
        pages=[PageText(page_no=1, text=full_text, char_count=len(full_text))],
        full_text=full_text,
        language=lang,
        language_confidence=conf,
    )
    return extracted, headings


def extract_marker_sections(path: str | Path) -> list[dict]:
    """Şablondaki açık bölüm işaretlerini okur (varsa).

    "Bölüm Başlangıcı – IX.Uçuş Kontrol Bilgisayarı" / "Bölüm Sonu – IX...."
    çiftlerini eşleştirip her bölüm için başlık + karakter aralığı döner.
    Bu işaretler varsa başlık tespitine hiç gerek kalmıyor: hangi bölümlerin
    zorunlu olduğunu ve nerede başlayıp bittiğini şablon zaten söylüyor.

    Dönen her öğe: {"numbering", "title", "char_start", "char_end"}
    İşaret yoksa boş liste döner (çağıran taraf sezgisel yola düşer).
    """
    import docx

    document = docx.Document(str(Path(path)))
    opens: dict[str, tuple[str, int]] = {}
    out: list[dict] = []
    cursor = 0

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        m = SECTION_MARKER_RE.match(text)
        if m:
            kind = m.group(1).lower()
            numbering = m.group(2).upper()
            title = m.group(3).strip(" .\t")
            if kind.startswith("bas") or kind.startswith("baş"):
                opens[numbering] = (title, cursor + len(text) + 2)
            else:
                start_title, start_off = opens.pop(numbering, (title, cursor))
                out.append({
                    "numbering": numbering,
                    "title": start_title or title,
                    "char_start": start_off,
                    "char_end": cursor,
                })
        cursor += len(text) + 2

    # Kapanmayan işaretler (şablon hatası) — dokümanın sonuna kadar say
    for numbering, (title, off) in opens.items():
        out.append({"numbering": numbering, "title": title,
                    "char_start": off, "char_end": cursor})

    out.sort(key=lambda x: x["char_start"])
    return out


def detect_section_level(headings: list[Heading]) -> int:
    """Ana bölümlerin hangi başlık seviyesinde olduğunu veriden tespit eder.

    Neden gerekli: şablonlar tutarlı değil. Jet Motor DTR'de Heading 1
    önsöz için kullanılmış (ŞEKİLLER, TABLOLAR, SİMGELER DİZİNİ), puanlı
    ana bölümler Heading 2'de. Savaşan İHA'da ana bölümler Heading 1'de.
    Sabit bir seviye varsaymak birinde tüm bölümleri kaçırıyor.

    Kural: puan ağırlığı taşıyan başlıkların en çok bulunduğu seviye ana
    bölüm seviyesidir (şablonlar puanı ana bölüme yazıyor). Puan hiç yoksa,
    en az 3 başlık içeren en küçük seviyeye düşülür.
    """
    style_heads = [h for h in headings if h.detected_by == "style" and h.level > 0]
    pool = style_heads or [h for h in headings if h.level > 0]
    if not pool:
        return 1

    scored: dict[int, int] = {}
    for h in pool:
        if h.points is not None:
            scored[h.level] = scored.get(h.level, 0) + 1
    if scored:
        best = max(scored.values())
        return min(lvl for lvl, n in scored.items() if n == best)

    counts: dict[int, int] = {}
    for h in pool:
        counts[h.level] = counts.get(h.level, 0) + 1
    for lvl in sorted(counts):
        if counts[lvl] >= 3:
            return lvl
    return min(counts)
