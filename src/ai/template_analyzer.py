"""T-Sistem · Rapor sablonu analiz motoru (GERCEK LLM).

ONCEKI DURUM
------------
* Hicbir LLM import edilmiyordu; tek bir regex ile numarali baslik araniyordu.
* Puan bulunamazsa her kritere 20.0 atanıyordu.
* `if len(rubric_items) < 3:` -> sablondan 3'ten az baslik yakalanirsa TUM cikti
  atilip asama koduna gore ELLE YAZILMIS sabit rubrik donduruluyordu. Gercek
  TEKNOFEST sablonlarinda bu regex nadiren 3+ eslesme buldugu icin PRATIKTE
  neredeyse her zaman sabit rubrik doner, hangi yarisma olursa olsun.

Kullanicinin en cok tekrarladigi sikayetin kok nedeni buydu:
  "rapor sablonu ile 4. adimdaki kriterler uyusmuyor, kafasina gore degil ait
   oldugu asamanin rapor sablonundaki puanlamaya gore yapilmali"

YENI DURUM
----------
* Sablonun metni VE TABLOLARI (puanlama tablolari genelde tablodadir) cikarilir.
* Gercek LLM, kriter + ALT KRITER + puan agirliklarini cikarir.
* Her kriter icin `source_quote` zorunlu; dogrulanamayan kriter elenir.
* Toplam 100'e normalize edilir ama ORAN KORUNUR.
* LLM yoksa sabit rubrik URETILMEZ; `LLMUnavailable` firlatilir.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from src.data.models import RubricCriterion

from .llm import LLMUnavailable, get_llm
from .spec_analyzer import MAX_CHARS, _normalize, verify_quote

log = logging.getLogger("tsistem.ai.template")

SYSTEM_PROMPT = (
    "Sen TEKNOFEST rapor sablonlarini inceleyen bir degerlendirme rubrigi uzmanisin. "
    "Gorevin, sablonda yazan puanlama tablosunu BIREBIR cikarmaktir. "
    "ASLA kendi kriterini uydurma, puan dagitimini kendi kafana gore degistirme. "
    "Yalnizca gecerli JSON dondur."
)

USER_TEMPLATE = """Asagida bir TEKNOFEST yarismasinin resmi RAPOR SABLONU yer aliyor.

YARISMA : {competition_name}
ASAMA   : {stage_code} ({stage_name})
{level_line}{branch_line}
SABLON ICERIGI (metin + tablolar):
\"\"\"
{text}
\"\"\"

GOREV
Bu sablonda tanimli DEGERLENDIRME KRITERLERINI ve PUAN AGIRLIKLARINI cikar.

KATI KURALLAR
1. Puanlar sablonda yaziyorsa BIREBIR onlari kullan. Kendi puanini uydurma.
2. Sablonda alt kriter varsa (or. "Algoritmalar 30 puan" icinde
   "Veri Setleri 10 / Algoritmalar 15 / Akis Semasi 5") bunlari `parent_code`
   ile bagla. Ust kriterin puani alt kriterlerin TOPLAMI olmalidir.
3. Her kriter icin `source_quote` alanina sablondan BIREBIR alinti koy.
   Dayanak bulamadigin kriteri EKLEME.
4. Sablonda puanlama tablosu YOKSA `criteria` dizisini BOS dondur ve
   `has_scoring_table` alanini false yap. SAKIN varsayilan kriter uydurma.
5. Ayrica sablonun zorunlu bolum basliklarini, sayfa sinirini ve yazi tipi /
   marj kurallarini da cikar (varsa).

JSON SEMASI (baska hicbir sey yazma):
{{
  "has_scoring_table": true,
  "criteria": [
    {{
      "criterion_code": "C1",
      "criterion_name": "Problem Tanimi ve Mevcut Durum",
      "description": "kriterin sablondaki aciklamasi",
      "max_score": 10,
      "parent_code": null,
      "source_quote": "sablondan birebir alinti"
    }}
  ],
  "required_sections": ["Ozet", "Giris", "Yontem"],
  "max_pages": null,
  "font_and_margins": null,
  "total_declared_score": 100,
  "summary": "sablonun 2 cumlelik ozeti"
}}"""


@dataclass
class TemplateAnalysis:
    criteria: list[RubricCriterion]
    required_sections: list[str]
    max_pages: int | None
    font_and_margins: str | None
    total_declared: float
    has_scoring_table: bool
    summary: str
    provider: str
    model: str
    dropped_unverified: int = 0
    warnings: list[str] = None  # type: ignore[assignment]

    def __post_init__(self) -> None:
        if self.warnings is None:
            self.warnings = []

    @property
    def total_score(self) -> float:
        return round(sum(c.max_score for c in self.criteria if not c.parent_code), 2)


# ── metin + tablo cikarimi ─────────────────────────────────────────────────
def extract_template_text(source: str | bytes | Path, *, max_pages: int = 40) -> str:
    """DOCX veya PDF'ten metin + TABLO hucrelerini cikarir.

    Puanlama tablolari genelde Word tablosundadir; eski kod tablo hucrelerini
    kismen okuyor ama regex'e takilmadigi icin kullanamiyordu.
    """
    if isinstance(source, str) and not Path(source).exists():
        return source[:MAX_CHARS]

    path = Path(source) if not isinstance(source, bytes) else None
    data = source if isinstance(source, bytes) else path.read_bytes()  # type: ignore[union-attr]

    if data[:2] == b"PK":  # DOCX (zip)
        return _extract_docx(data)[:MAX_CHARS]
    if data[:4] == b"%PDF":
        return _extract_pdf(data, max_pages)[:MAX_CHARS]
    return data.decode("utf-8", errors="replace")[:MAX_CHARS]


def _extract_docx(data: bytes) -> str:
    try:
        import io

        import docx  # type: ignore
    except ImportError as exc:
        raise RuntimeError("DOCX okumak icin python-docx gerekli: pip install python-docx") from exc

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for t_index, table in enumerate(document.tables, 1):
        parts.append(f"\n[TABLO {t_index}]")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            # ardisik ayni hucreleri (birlestirilmis hucre) tekilleştir
            deduped: list[str] = []
            for cell in cells:
                if not deduped or deduped[-1] != cell:
                    deduped.append(cell)
            if any(deduped):
                parts.append(" | ".join(deduped))
    text = "\n".join(parts)
    if not text.strip():
        raise RuntimeError("DOCX icinden metin cikarilamadi.")
    return text


def _extract_pdf(data: bytes, max_pages: int) -> str:
    try:
        import pymupdf  # type: ignore
    except ImportError:
        try:
            import fitz as pymupdf  # type: ignore
        except ImportError as exc:
            raise RuntimeError("PDF okumak icin pymupdf gerekli.") from exc

    parts: list[str] = []
    with pymupdf.open(stream=data, filetype="pdf") as doc:
        for page_no, page in enumerate(doc, 1):
            if page_no > max_pages:
                break
            parts.append(page.get_text())
            try:
                for table in page.find_tables().tables:
                    for row in table.extract():
                        cells = [str(c or "").strip() for c in row]
                        if any(cells):
                            parts.append(" | ".join(cells))
            except Exception as exc:  # noqa: BLE001 - tablo bulucu opsiyonel
                log.debug("[template] tablo cikarimi atlandi (sayfa %d): %s", page_no, exc)
    text = "\n".join(parts)
    if not text.strip():
        raise RuntimeError("PDF sablonundan metin cikarilamadi (taranmis olabilir).")
    return text


# ── ana giris ──────────────────────────────────────────────────────────────
def analyze_template(
    source: str | bytes | Path,
    *,
    competition_id: str,
    competition_name: str,
    stage_code: str,
    stage_name: str = "",
    level: str = "Genel",
    branch_code: str | None = None,
    normalize_to: float | None = 100.0,
) -> TemplateAnalysis:
    """Sablondan rubrigi cikarir. LLM yoksa `LLMUnavailable` firlatir."""
    text = extract_template_text(source)
    if len(text.strip()) < 300:
        raise RuntimeError("Sablon metni analiz icin fazla kisa.")

    prompt = USER_TEMPLATE.format(
        competition_name=competition_name,
        stage_code=stage_code.upper(),
        stage_name=stage_name or stage_code.upper(),
        level_line=f"SEVIYE  : {level}\n" if level and level != "Genel" else "",
        branch_line=f"ALT DAL : {branch_code}\n" if branch_code else "",
        text=text,
    )

    payload, result = get_llm().complete_json(
        prompt, system=SYSTEM_PROMPT, max_tokens=6000, temperature=0.05,
        validator=_validate_payload,
    )

    haystack = _normalize(text)
    criteria: list[RubricCriterion] = []
    dropped = 0
    warnings: list[str] = []

    for idx, raw in enumerate(payload.get("criteria", [])):
        quote = str(raw.get("source_quote") or "").strip()
        if not quote or not verify_quote(quote, haystack):
            dropped += 1
            continue
        try:
            max_score = float(raw.get("max_score") or 0)
        except (TypeError, ValueError):
            dropped += 1
            continue
        if max_score <= 0:
            dropped += 1
            continue
        criteria.append(
            RubricCriterion(
                competition_id=competition_id,
                stage_code=stage_code.upper(),
                level=level,
                branch_code=branch_code,
                criterion_code=str(raw.get("criterion_code") or f"C{idx + 1}")[:16],
                criterion_name=str(raw.get("criterion_name") or f"Kriter {idx + 1}")[:160],
                description=(str(raw.get("description")) if raw.get("description") else None),
                max_score=max_score,
                parent_code=(str(raw["parent_code"]) if raw.get("parent_code") else None),
                source_quote=quote,
                order_index=idx,
            )
        )

    has_table = bool(payload.get("has_scoring_table")) and bool(criteria)

    if dropped:
        warnings.append(f"{dropped} kriter, sablonda dayanak bulunamadigi icin elendi.")

    if not criteria:
        warnings.append(
            "Bu sablonda dogrulanabilir bir puanlama tablosu bulunamadi. "
            "Kriterleri yonetici panelinden elle tanimlamaniz gerekiyor — "
            "sistem varsayilan bir rubrik URETMEZ."
        )
    else:
        _validate_parents(criteria, warnings)
        if normalize_to:
            _normalize_scores(criteria, normalize_to, warnings)

    log.info(
        "[template] %s/%s (%s) -> %d kriter (%d elendi) · toplam %.1f · %s/%s",
        competition_id, stage_code, level, len(criteria), dropped,
        sum(c.max_score for c in criteria if not c.parent_code),
        result.provider, result.model,
    )

    return TemplateAnalysis(
        criteria=criteria,
        required_sections=[str(s) for s in (payload.get("required_sections") or []) if s],
        max_pages=_int_or_none(payload.get("max_pages")),
        font_and_margins=(str(payload["font_and_margins"]) if payload.get("font_and_margins") else None),
        total_declared=float(payload.get("total_declared_score") or 0) or 0.0,
        has_scoring_table=has_table,
        summary=str(payload.get("summary") or "").strip(),
        provider=result.provider,
        model=result.model,
        dropped_unverified=dropped,
        warnings=warnings,
    )


def _validate_parents(criteria: list[RubricCriterion], warnings: list[str]) -> None:
    """Ust kriterin puani alt kriterlerin toplamina esit olmali."""
    by_code = {c.criterion_code: c for c in criteria}
    children: dict[str, list[RubricCriterion]] = {}
    for crit in criteria:
        if crit.parent_code:
            if crit.parent_code not in by_code:
                crit.parent_code = None  # yetim alt kriter -> ust kritere terfi
                continue
            children.setdefault(crit.parent_code, []).append(crit)

    for parent_code, kids in children.items():
        parent = by_code[parent_code]
        child_total = round(sum(k.max_score for k in kids), 2)
        if abs(child_total - parent.max_score) > 0.51:
            warnings.append(
                f"'{parent.criterion_name}' puani {parent.max_score:g}, alt kriterleri "
                f"{child_total:g} ediyor. Ust kriter alt toplama esitlendi."
            )
            parent.max_score = child_total


def _normalize_scores(criteria: list[RubricCriterion], target: float, warnings: list[str]) -> None:
    """Toplam hedeften saparsa ORANI koruyarak olceklendirir."""
    tops = [c for c in criteria if not c.parent_code]
    total = round(sum(c.max_score for c in tops), 2)
    if total <= 0 or abs(total - target) < 0.01:
        return

    factor = target / total
    warnings.append(
        f"Sablondaki toplam puan {total:g}; {target:g} uzerinden oranlanarak olceklendirildi."
    )
    for crit in criteria:
        crit.max_score = round(crit.max_score * factor, 1)

    # yuvarlama farkini en buyuk kritere ekle
    tops = [c for c in criteria if not c.parent_code]
    drift = round(target - sum(c.max_score for c in tops), 1)
    if abs(drift) >= 0.1 and tops:
        biggest = max(tops, key=lambda c: c.max_score)
        biggest.max_score = round(biggest.max_score + drift, 1)


def _validate_payload(payload: Any) -> dict[str, Any]:
    if not isinstance(payload, dict):
        raise ValueError("Kok nesne sozluk olmali.")
    crits = payload.get("criteria")
    if not isinstance(crits, list):
        raise ValueError("'criteria' bir dizi olmali.")
    for item in crits:
        if not isinstance(item, dict):
            raise ValueError("Her kriter bir nesne olmali.")
        if not item.get("criterion_name"):
            raise ValueError("Her kriterin 'criterion_name' alani olmali.")
        if "max_score" not in item:
            raise ValueError("Her kriterin 'max_score' alani olmali.")
        if "source_quote" not in item:
            raise ValueError("Her kriterin 'source_quote' alani olmali.")
    return payload


def _int_or_none(value: Any) -> int | None:
    if value in (None, "", "null"):
        return None
    try:
        return int(re.sub(r"[^\d]", "", str(value)) or 0) or None
    except (TypeError, ValueError):
        return None


__all__ = ["analyze_template", "TemplateAnalysis", "extract_template_text", "LLMUnavailable"]
